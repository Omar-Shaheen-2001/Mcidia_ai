from flask import Blueprint, render_template, request, flash, session, redirect, url_for, current_app, jsonify, send_file
from flask_jwt_extended import get_jwt_identity
from utils.decorators import login_required
from models import AILog, ChatSession, Service
from utils.ai_providers.ai_manager import AIManager
from datetime import datetime
from werkzeug.utils import secure_filename
from weasyprint import HTML
from io import BytesIO
import json
import time
import os
import re

consultation_bp = Blueprint('consultation', __name__)

def format_response_html(text):
    """Convert plain text response to professional HTML (Python equivalent of formatResponse JS)"""
    html = ''
    
    # Split into paragraphs
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    for idx, para in enumerate(paragraphs):
        # Check if it's a table (HTML table)
        if '<table' in para or '<TR>' in para or '<tr>' in para:
            html += para.replace('<table', '<table class="response-table"')
        # Check if it's a markdown-style table
        elif '|' in para and len(para.split('\n')) >= 2 and '-' in para.split('\n')[1]:
            lines = para.split('\n')
            table_html = '<table class="response-table"><thead><tr>'
            
            # Parse header
            headers = [h.strip() for h in lines[0].split('|') if h.strip()]
            for h in headers:
                table_html += f'<th>{h}</th>'
            table_html += '</tr></thead><tbody>'
            
            # Parse rows (skip separator line)
            for i in range(2, len(lines)):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                if cells:
                    table_html += '<tr>'
                    for cell in cells:
                        table_html += f'<td>{cell}</td>'
                    table_html += '</tr>'
            table_html += '</tbody></table>'
            html += table_html
        # Check if it's a numbered list
        elif re.match(r'^\d+\.', para):
            html += '<div class="response-section">'
            items = [l.strip() for l in para.split('\n') if l.strip()]
            for item in items:
                match = re.match(r'^\d+\.\s*(.+)', item)
                if match:
                    html += f'<div class="response-text"><strong>{match.group(1)}</strong></div>'
            html += '</div>'
        # Check if it's a bullet list
        elif re.match(r'^[-•*]', para):
            html += '<ul class="response-list">'
            items = [l.strip() for l in para.split('\n') if l.strip()]
            for item in items:
                text_content = re.sub(r'^[-•*]\s*', '', item).strip()
                if text_content:
                    html += f'<li>{text_content}</li>'
            html += '</ul>'
        # Check if it's a title/header (first paragraph or short text)
        elif idx == 0 or len(para) < 100:
            html += f'<div class="response-title">{para}</div>'
        # Regular paragraph
        else:
            html += f'<div class="response-text">{para}</div>'
    
    return html or text

@consultation_bp.route('/')
def index():
    lang = session.get('language', 'ar')
    db = current_app.extensions['sqlalchemy']
    
    # Get available services for consultation topics and convert to dict
    services = db.session.query(Service).filter_by(is_active=True).all()
    services_list = [{'id': s.id, 'title_ar': s.title_ar, 'title_en': s.title_en, 'slug': s.slug} for s in services]
    
    # Get recent sessions for current user if logged in
    recent_sessions = []
    try:
        from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity
        verify_jwt_in_request(optional=True, locations=['cookies'])
        user_id = get_jwt_identity()
        if user_id:
            recent_sessions = db.session.query(ChatSession).filter_by(
                user_id=int(user_id)
            ).order_by(ChatSession.updated_at.desc()).limit(5).all()
    except:
        pass
    
    return render_template('consultation/index.html', lang=lang, services=services_list, recent_sessions=recent_sessions)

@consultation_bp.route('/api/sessions', methods=['GET'])
@login_required
def get_sessions_api():
    """API endpoint to get user's consultation sessions"""
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    
    # Get paginated sessions
    sessions = db.session.query(ChatSession).filter_by(
        user_id=user_id
    ).order_by(ChatSession.updated_at.desc()).paginate(page=page, per_page=per_page)
    
    sessions_list = []
    for s in sessions.items:
        try:
            messages = json.loads(s.messages) if s.messages else []
            message_count = len(messages)
            total_cost = sum(m.get('cost', 0) for m in messages if m.get('role') == 'assistant')
        except:
            message_count = 0
            total_cost = 0
        
        sessions_list.append({
            'id': s.id,
            'domain': s.domain,
            'message_count': message_count,
            'total_cost': round(total_cost, 4),
            'created_at': s.created_at.isoformat(),
            'updated_at': s.updated_at.isoformat()
        })
    
    return jsonify({
        'sessions': sessions_list,
        'total': sessions.total,
        'pages': sessions.pages,
        'current_page': page
    })

@consultation_bp.route('/sessions')
@login_required
def all_sessions():
    """View all consultation sessions"""
    lang = session.get('language', 'ar')
    return render_template('consultation/all_sessions.html', lang=lang)

@consultation_bp.route('/start')
@login_required
def start_session():
    """Start a new consultation session"""
    lang = session.get('language', 'ar')
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    
    # Get available services for topics and convert to dict
    services = db.session.query(Service).filter_by(is_active=True).all()
    services_list = [{'id': s.id, 'title_ar': s.title_ar, 'title_en': s.title_en, 'slug': s.slug} for s in services]
    
    return render_template('consultation/start.html', lang=lang, services=services_list)

@consultation_bp.route('/session/<int:session_id>')
@login_required
def view_session(session_id):
    """View an existing consultation session"""
    lang = session.get('language', 'ar')
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    
    # Get the session
    chat_session = db.session.query(ChatSession).filter_by(
        id=session_id,
        user_id=user_id
    ).first()
    
    if not chat_session:
        flash('جلسة غير موجودة' if lang == 'ar' else 'Session not found', 'danger')
        return redirect(url_for('consultation.index'))
    
    # Parse messages
    try:
        messages = json.loads(chat_session.messages) if chat_session.messages else []
    except:
        messages = []
    
    # Get available services and convert to dict
    services = db.session.query(Service).filter_by(is_active=True).all()
    services_list = [{'id': s.id, 'title_ar': s.title_ar, 'title_en': s.title_en, 'slug': s.slug} for s in services]
    
    return render_template(
        'consultation/session.html',
        lang=lang,
        session=chat_session,
        messages=messages,
        services=services_list
    )

@consultation_bp.route('/api/send-message', methods=['POST'])
@login_required
def send_message():
    """API endpoint to send a message in consultation"""
    start_time = time.time()
    
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    data = request.get_json()
    
    message = data.get('message')
    session_id = data.get('session_id')
    topic = data.get('topic', 'General Consultation')
    
    if not message:
        return jsonify({'error': 'Message is required'}), 400
    
    try:
        # RAG integration: Get relevant context from vector store
        rag_context = ""
        try:
            from utils.knowledge.embeddings import create_embedding
            from utils.knowledge.vector_store import get_vector_store
            from models import User
            
            user = db.session.query(User).filter_by(id=user_id).first()
            if user and user.organization_id:
                query_embedding = create_embedding(message)
                if query_embedding:
                    vector_store = get_vector_store()
                    context_docs = vector_store.search(query_embedding, user.organization_id, top_k=3)
                    
                    if context_docs:
                        rag_context = "\n**السياق من قاعدة المعرفة:**\n"
                        for doc in context_docs:
                            rag_context += f"- {doc['text'][:200]}...\n"
        except:
            pass
        
        # Call AI using AIManager (same as other consulting modules)
        ai = AIManager.for_use_case('consultation')
        
        # Get chart generation instructions
        try:
            from utils.chart_generator import get_ai_chart_instructions
            chart_instructions = get_ai_chart_instructions('ar' if any(ord(c) > 127 for c in message) else 'en')
        except:
            chart_instructions = ""
        
        system_prompt = f"""أنت مستشار خبير متخصص في مجال {topic}.
تقدّم استشارات عملية وقيّمة وقابلة للتطبيق.
كن موجزاً وفعالاً في إجابتك.
الرد بالعربية إذا كانت الأسئلة بالعربية، والإنجليزية إذا كانت بالإنجليزية.

**أهم: أضف مخططات بصرية لتوضيح إجاباتك:**
- إذا تحدثت عن أرقام أو إحصائيات، استخدم مخطط!
- إذا قارنت بين عناصر متعددة، استخدم مخطط!
- إذا تحدثت عن نسب أو توزيعات، استخدم مخطط دائري!
- إذا تحدثت عن اتجاهات زمنية، استخدم مخطط خطي!

{chart_instructions}"""
        
        augmented_message = f"{rag_context}\n\n**السؤال:** {message}" if rag_context else message
        ai_response = ai.chat(system_prompt, augmented_message)
        execution_time_ms = int((time.time() - start_time) * 1000)
        
        # Extract charts from AI response
        charts = []
        cleaned_response = ai_response
        try:
            from utils.chart_generator import process_ai_response_for_charts
            cleaned_response, charts = process_ai_response_for_charts(ai_response)
        except:
            pass
        
        # Estimate tokens and cost (rough estimation)
        estimated_tokens = len(message.split()) + len(ai_response.split())
        # OpenAI pricing: $0.0015 per 1K input, $0.002 per 1K output
        estimated_cost = (estimated_tokens * 0.00175 / 1000) if estimated_tokens > 0 else 0.001
        
        # Get provider info
        available_providers = AIManager.get_available_providers()
        provider_name = available_providers[0] if available_providers else 'huggingface'
        provider_config = AIManager.get_provider_info(provider_name)
        model_name = provider_config.get('default_model', 'llama3')
        
        # Log AI usage
        log = AILog(
            user_id=user_id,
            module='consultation',
            service_type=topic,
            provider_type=provider_name,
            model_name=model_name,
            prompt=message,
            response=ai_response,
            tokens_used=estimated_tokens,
            estimated_cost=estimated_cost,
            execution_time_ms=execution_time_ms,
            status='success'
        )
        db.session.add(log)
        
        # Update or create chat session
        if session_id:
            chat_session = db.session.query(ChatSession).filter_by(
                id=session_id,
                user_id=user_id
            ).first()
        else:
            chat_session = ChatSession(user_id=user_id, domain=topic)
        
        # Parse existing messages
        try:
            messages = json.loads(chat_session.messages) if chat_session.messages else []
        except:
            messages = []
        
        # Add new messages
        messages.append({
            'role': 'user',
            'content': message,
            'timestamp': datetime.utcnow().isoformat()
        })
        messages.append({
            'role': 'assistant',
            'content': cleaned_response,
            'timestamp': datetime.utcnow().isoformat(),
            'cost': estimated_cost,
            'charts': charts if charts else None
        })
        
        chat_session.messages = json.dumps(messages)
        chat_session.updated_at = datetime.utcnow()
        
        if not session_id:
            db.session.add(chat_session)
        
        db.session.commit()
        
        return jsonify({
            'response': cleaned_response,
            'charts': charts if charts else [],
            'session_id': chat_session.id,
            'cost': estimated_cost,
            'timestamp': datetime.utcnow().isoformat()
        })
        
    except Exception as e:
        from openai import APIError, APIConnectionError, RateLimitError, AuthenticationError
        
        execution_time_ms = int((time.time() - start_time) * 1000)
        error_str = str(e)
        
        # Handle specific OpenAI API errors
        error_message = error_str
        http_status = 500
        
        if isinstance(e, AuthenticationError):
            error_message = 'خطأ في المصادقة - تحقق من API Key / Authentication error - check your API Key'
            http_status = 401
        elif isinstance(e, RateLimitError):
            error_message = 'تم تجاوز حد الاستخدام - يرجى المحاولة لاحقاً / Rate limit exceeded - please try again later'
            http_status = 429
        elif 'insufficient_quota' in error_str.lower():
            error_message = 'حد الاستخدام قد تم تجاوزه - يرجى التحقق من حسابك في OpenAI وإضافة رصيد / Insufficient quota - please check your OpenAI account and add credits'
            http_status = 402
        elif isinstance(e, APIConnectionError):
            error_message = 'خطأ في الاتصال - تحقق من الإنترنت / Connection error - check your internet'
            http_status = 503
        elif isinstance(e, APIError):
            error_message = f'خطأ في API: {error_str} / API Error: {error_str}'
            http_status = 500
        
        # Log failed attempt
        failed_log = AILog(
            user_id=user_id,
            module='consultation',
            service_type=topic,
            provider_type='openai',
            model_name='gpt-3.5-turbo',
            prompt=message,
            response='',
            status='failed',
            error_message=error_str,
            execution_time_ms=execution_time_ms
        )
        db.session.add(failed_log)
        db.session.commit()
        
        return jsonify({'error': error_message}), http_status

@consultation_bp.route('/api/create-session', methods=['POST'])
@login_required
def create_session():
    """API endpoint to create a new consultation session"""
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    data = request.get_json()
    
    topic = data.get('topic', 'General Consultation')
    
    # Check if service has file upload enabled
    enable_file_upload = True  # Default to True - allow file uploads by default
    
    try:
        # Find service by title (try both Arabic and English)
        service = db.session.query(Service).filter_by(title_ar=topic).first()
        if not service:
            service = db.session.query(Service).filter_by(title_en=topic).first()
        
        # Debug log
        current_app.logger.info(f"Service search for topic: '{topic}' - Found: {service is not None}")
        
        # If service found and has offerings, check if any offering has file upload enabled
        if service and service.offerings and len(service.offerings) > 0:
            # If ANY offering has file upload enabled, show the button
            enable_file_upload = any(o.enable_file_upload for o in service.offerings)
            current_app.logger.info(f"Service '{topic}' has {len(service.offerings)} offerings. Any enabled: {enable_file_upload}")
            current_app.logger.info(f"Offerings details: {[(o.title_ar, o.enable_file_upload) for o in service.offerings]}")
        else:
            current_app.logger.info(f"Service '{topic}' not found or has no offerings. Default: file upload enabled")
            enable_file_upload = True  # Default to True
    except Exception as e:
        current_app.logger.error(f"Error checking service file upload: {str(e)}")
        enable_file_upload = True  # Default to True on error
    
    # Create new session
    chat_session = ChatSession(
        user_id=user_id,
        domain=topic,
        messages=json.dumps([]),
        enable_file_upload=enable_file_upload
    )
    db.session.add(chat_session)
    db.session.commit()
    
    return jsonify({
        'session_id': chat_session.id,
        'domain': chat_session.domain,
        'enable_file_upload': enable_file_upload,
        'created_at': chat_session.created_at.isoformat()
    })

# ==================== DOCUMENT UPLOAD FOR AI ANALYSIS ====================

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        current_app.logger.error(f"PDF extraction error: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from Word document using python-docx"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        current_app.logger.error(f"DOCX extraction error: {str(e)}")
        return None

def extract_text_from_txt(file_path):
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()
    except Exception as e:
        current_app.logger.error(f"TXT extraction error: {str(e)}")
        return None

@consultation_bp.route('/api/upload-document', methods=['POST'])
@login_required
def upload_document():
    """API endpoint to upload and extract text from documents for AI analysis"""
    user_id = int(get_jwt_identity())
    
    # Check if file was uploaded
    if 'document' not in request.files:
        return jsonify({'success': False, 'error': 'لم يتم تحميل ملف / No file uploaded'}), 400
    
    file = request.files['document']
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'لم يتم اختيار ملف / No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'نوع الملف غير مدعوم. يرجى استخدام PDF, Word, أو TXT / File type not supported. Please use PDF, Word, or TXT'}), 400
    
    # Check file size
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    
    if file_size > MAX_FILE_SIZE:
        return jsonify({'success': False, 'error': 'حجم الملف كبير جداً (أقصى 10MB) / File too large (max 10MB)'}), 400
    
    try:
        # Secure filename and save
        filename = secure_filename(file.filename)
        timestamp = int(time.time())
        unique_filename = f"{user_id}_{timestamp}_{filename}"
        
        # Create upload directory if not exists
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'documents', 'consultation')
        os.makedirs(upload_dir, exist_ok=True)
        
        file_path = os.path.join(upload_dir, unique_filename)
        file.save(file_path)
        
        # Extract text based on file type
        extension = filename.rsplit('.', 1)[1].lower()
        extracted_text = None
        
        if extension == 'pdf':
            extracted_text = extract_text_from_pdf(file_path)
        elif extension in ['doc', 'docx']:
            extracted_text = extract_text_from_docx(file_path)
        elif extension == 'txt':
            extracted_text = extract_text_from_txt(file_path)
        
        # Clean up file after extraction
        try:
            os.remove(file_path)
        except:
            pass
        
        if not extracted_text:
            return jsonify({
                'success': False,
                'error': 'فشل استخراج النص من الملف / Failed to extract text from file'
            }), 500
        
        # Truncate if too long (max 15000 characters for context)
        max_length = 15000
        if len(extracted_text) > max_length:
            extracted_text = extracted_text[:max_length] + "\n... [تم اقتطاع النص بسبب الطول / Text truncated due to length]"
        
        # Create preview (first 500 characters)
        preview = extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text
        
        return jsonify({
            'success': True,
            'filename': filename,
            'content': extracted_text,
            'preview': preview,
            'length': len(extracted_text)
        })
        
    except Exception as e:
        current_app.logger.error(f"Document upload error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'خطأ في معالجة الملف / Error processing file: {str(e)}'
        }), 500

# ==================== STRATEGIC PLANNING SUBMODULE ====================

@consultation_bp.route('/export-pdf/<int:session_id>', methods=['GET', 'POST'])
@login_required
def export_session_pdf(session_id):
    """Export consultation session as professional PDF"""
    lang = session.get('language', 'ar')
    db = current_app.extensions['sqlalchemy']
    user_id = int(get_jwt_identity())
    
    # Get the session
    chat_session = db.session.query(ChatSession).filter_by(
        id=session_id,
        user_id=user_id
    ).first()
    
    if not chat_session:
        return jsonify({'error': 'Session not found'}), 404
    
    # Parse messages
    try:
        messages = json.loads(chat_session.messages) if chat_session.messages else []
    except:
        messages = []
    
    # Get chart images from request if POST
    chart_images = {}
    if request.method == 'POST':
        data = request.get_json() or {}
        chart_images = data.get('chartImages', {})
    
    # Calculate total cost
    total_cost = sum(m.get('cost', 0) for m in messages if m.get('role') == 'assistant')
    
    # Generate HTML for PDF with proper styling
    html_content = f"""
    <!DOCTYPE html>
    <html dir="{'rtl' if lang == 'ar' else 'ltr'}" lang="{lang}">
    <head>
        <meta charset="UTF-8">
        <link href="https://fonts.googleapis.com/css2?family=Cairo:wght@400;500;600;700&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
        <style>
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            
            @page {{
                size: A4;
                margin: 2cm;
            }}
            
            body {{
                font-family: {'Cairo' if lang == 'ar' else 'Inter'}, Arial, sans-serif;
                direction: {'rtl' if lang == 'ar' else 'ltr'};
                text-align: {'right' if lang == 'ar' else 'left'};
                line-height: 1.8;
                color: #333;
                font-size: 14px;
                background: white;
            }}
            
            .header {{
                background: linear-gradient(135deg, #0A2756 0%, #2767B1 100%);
                color: white;
                padding: 30px;
                margin-bottom: 30px;
                border-radius: 8px;
                text-align: center;
                page-break-inside: avoid;
            }}
            
            .header h1 {{
                font-size: 28px;
                font-weight: 700;
                margin-bottom: 15px;
            }}
            
            .header-meta {{
                font-size: 13px;
                opacity: 0.95;
                line-height: 1.8;
            }}
            
            .header-meta strong {{
                font-weight: 600;
            }}
            
            .message {{
                margin-bottom: 25px;
                padding: 18px;
                border-{'right' if lang == 'ar' else 'left'}: 5px solid #0A2756;
                background: #f8f9fa;
                border-radius: 4px;
                page-break-inside: avoid;
            }}
            
            .message.assistant {{
                border-{'right' if lang == 'ar' else 'left'}-color: #2C8C56;
                background: #f0f8f5;
            }}
            
            .message-role {{
                font-weight: 700;
                font-size: 12px;
                text-transform: uppercase;
                color: #0A2756;
                margin-bottom: 10px;
                letter-spacing: 0.5px;
            }}
            
            .message.assistant .message-role {{
                color: #2C8C56;
            }}
            
            .message-content {{
                font-size: 14px;
                line-height: 1.8;
                color: #333;
                word-wrap: break-word;
                overflow-wrap: break-word;
            }}
            
            .response-table {{
                width: 100%;
                border-collapse: collapse;
                margin: 15px 0;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 2px 8px rgba(10, 39, 86, 0.08);
                page-break-inside: avoid;
            }}
            
            .response-table thead {{
                background: #0A2756 !important;
            }}
            
            .response-table thead th {{
                padding: 14px 12px;
                text-align: {'right' if lang == 'ar' else 'left'};
                font-weight: 600;
                font-size: 13px;
                letter-spacing: 0.3px;
                border: none;
                color: #FFFFFF !important;
                background: #0A2756 !important;
            }}
            
            .response-table tbody td {{
                padding: 12px;
                border-bottom: 1px solid #f0f0f0;
                font-size: 13px;
                color: #444;
                text-align: {'right' if lang == 'ar' else 'left'};
            }}
            
            .response-table tbody tr:last-child td {{
                border-bottom: none;
            }}
            
            .response-table tbody tr:nth-child(even) {{
                background: #fafbfc;
            }}
            
            .response-table tbody tr:nth-child(odd) {{
                background: #ffffff;
            }}
            
            .message-content table {{
                width: 100%;
                border-collapse: collapse;
                margin: 15px 0;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 2px 8px rgba(10, 39, 86, 0.08);
                page-break-inside: avoid;
            }}
            
            .message-content table thead {{
                background: #0A2756 !important;
            }}
            
            .message-content table th {{
                padding: 14px 12px;
                text-align: {'right' if lang == 'ar' else 'left'};
                font-weight: 600;
                font-size: 13px;
                letter-spacing: 0.3px;
                border: none;
                color: #FFFFFF !important;
                background: #0A2756 !important;
            }}
            
            .message-content table td {{
                padding: 12px;
                border-bottom: 1px solid #f0f0f0;
                font-size: 13px;
                color: #444;
                text-align: {'right' if lang == 'ar' else 'left'};
            }}
            
            .message-content table tr:last-child td {{
                border-bottom: none;
            }}
            
            .message-content table tr:nth-child(even) {{
                background: #fafbfc;
            }}
            
            .message-content table tr:nth-child(odd) {{
                background: #ffffff;
            }}
            
            .chart-image {{
                width: 100%;
                height: auto;
                margin: 15px 0;
                text-align: center;
                page-break-inside: avoid;
            }}
            
            .chart-image img {{
                max-width: 100%;
                height: auto;
                border: 1px solid #e8ecf0;
                border-radius: 4px;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            }}
            
            .message-cost {{
                font-size: 12px;
                color: #FFC107;
                margin-top: 12px;
                font-weight: 600;
                text-align: {'left' if lang == 'ar' else 'right'};
            }}
            
            .total-cost {{
                text-align: {'left' if lang == 'ar' else 'right'};
                font-size: 16px;
                font-weight: 700;
                color: #0A2756;
                margin-top: 30px;
                margin-bottom: 30px;
                padding: 15px;
                background: #f0f8f5;
                border-radius: 4px;
                page-break-inside: avoid;
            }}
            
            .footer {{
                margin-top: 40px;
                padding-top: 20px;
                border-top: 2px solid #0A2756;
                text-align: center;
                font-size: 12px;
                color: #666;
                page-break-inside: avoid;
            }}
            
            .response-section, .response-title, .response-text, .response-list {{
                margin-bottom: 10px;
            }}
            
            .response-title {{
                font-size: 15px;
                font-weight: 700;
                color: #0A2756;
                margin-bottom: 8px;
            }}
            
            .response-text {{
                font-size: 14px;
                line-height: 1.8;
                margin-bottom: 8px;
            }}
            
            .response-list {{
                padding-{'right' if lang == 'ar' else 'left'}: 20px;
            }}
            
            .response-list li {{
                margin-bottom: 6px;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{"تقرير الاستشارة" if lang == 'ar' else 'Consultation Report'}</h1>
            <div class="header-meta">
                <div><strong>{"الموضوع:" if lang == 'ar' else 'Topic:'}</strong> {chat_session.domain}</div>
                <div><strong>{"رقم الجلسة:" if lang == 'ar' else 'Session #:'}</strong> {chat_session.id}</div>
                <div><strong>{"عدد الرسائل:" if lang == 'ar' else 'Messages:'}</strong> {len(messages)}</div>
                <div><strong>{"التاريخ:" if lang == 'ar' else 'Date:'}</strong> {datetime.now().strftime("%Y-%m-%d %H:%M")}</div>
            </div>
        </div>
        
        <div class="messages-container">
    """
    
    for idx, msg in enumerate(messages):
        content = msg.get('content', '')
        
        # Format assistant messages with professional HTML styling
        if msg.get('role') == 'assistant':
            # Use format_response_html to convert to professional HTML
            display_content = format_response_html(content)
        else:
            # For user messages, escape HTML
            display_content = (content.replace('&', '&amp;')
                                     .replace('<', '&lt;')
                                     .replace('>', '&gt;'))
            display_content = display_content.replace('\n', '<br>')
        
        cost_html = f"<div class='message-cost'>💰 ${msg.get('cost', 0):.4f}</div>" if msg.get('cost') else ""
        role_label = "المستخدم" if lang == 'ar' else "User" if msg.get('role') == 'user' else "المساعد" if lang == 'ar' else "Assistant"
        
        html_content += f"""
        <div class="message {msg.get('role', 'user')}">
            <div class="message-role">{role_label}</div>
            <div class="message-content">{display_content}</div>
            {cost_html}
        </div>
        """
        
        # Add chart image if available
        chart_key = f"chart-{idx}"
        if chart_key in chart_images:
            html_content += f"""
            <div class="chart-image">
                <img src="{chart_images[chart_key]}" alt="Chart" />
            </div>
            """
    
    html_content += f"""
        </div>
        
        <div class="total-cost">
            {"التكلفة الإجمالية:" if lang == 'ar' else 'Total Cost:'} ${total_cost:.4f}
        </div>
        
        <div class="footer">
            <p>{"تم إنشاء هذا التقرير بواسطة منصة Mcidia" if lang == 'ar' else 'Generated by Mcidia Platform'}</p>
            <p style="font-size: 11px; color: #999; margin-top: 8px;">{"هذا الملف يحتوي على معلومات سرية" if lang == 'ar' else 'This file contains confidential information'}</p>
        </div>
    </body>
    </html>
    """
    
    try:
        # Generate PDF using WeasyPrint
        pdf = HTML(string=html_content).write_pdf()
        
        # Create BytesIO object
        pdf_io = BytesIO(pdf)
        pdf_io.seek(0)
        
        filename = f"consultation_{chat_session.id}_{int(time.time())}.pdf"
        
        return send_file(
            pdf_io,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        current_app.logger.error(f"PDF export error: {str(e)}")
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500

@consultation_bp.route('/strategic-planning')
@login_required
def strategic_planning_module():
    """Strategic Planning consulting submodule - entry point"""
    lang = session.get('language', 'ar')
    return redirect(url_for('strategic_planning_ai.index'))

@consultation_bp.route('/strategic-planning/create', methods=['GET', 'POST'])
@login_required
def strategic_planning_create():
    """Create new strategic plan through consultation"""
    return redirect(url_for('strategic_planning_ai.create_plan'))
