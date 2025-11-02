from flask import Blueprint, render_template, request, flash, session, redirect, url_for
from flask_jwt_extended import get_jwt_identity
from utils.decorators import login_required
from werkzeug.utils import secure_filename
from models import User, Document, AILog
from app import db
from utils.ai_client import llm_chat
import os
import PyPDF2
from docx import Document as DocxDocument

knowledge_bp = Blueprint('knowledge', __name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(filepath, file_type):
    """Extract text from uploaded files"""
    try:
        if file_type == 'txt':
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_type == 'pdf':
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ''
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        elif file_type == 'docx':
            doc = DocxDocument(filepath)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        return ""

@knowledge_bp.route('/')
@login_required
def index():
    user_id = get_jwt_identity()
    documents = Document.query.filter_by(user_id=user_id).all()
    lang = session.get('language', 'ar')
    return render_template('knowledge/index.html', documents=documents, lang=lang)

@knowledge_bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    user_id = get_jwt_identity()
    lang = session.get('language', 'ar')
    
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('لم يتم اختيار ملف / No file selected', 'danger')
            return redirect(request.url)
        
        file = request.files['file']
        
        if file.filename == '':
            flash('لم يتم اختيار ملف / No file selected', 'danger')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_type = filename.rsplit('.', 1)[1].lower()
            filepath = os.path.join(UPLOAD_FOLDER, f"{user_id}_{filename}")
            file.save(filepath)
            
            # Extract text
            content_text = extract_text_from_file(filepath, file_type)
            
            # Save to database
            document = Document(
                user_id=user_id,
                filename=filename,
                file_type=file_type,
                file_path=filepath,
                content_text=content_text
            )
            db.session.add(document)
            db.session.commit()
            
            flash('تم رفع الملف بنجاح! / File uploaded successfully!', 'success')
            return redirect(url_for('knowledge.index'))
        else:
            flash('نوع الملف غير مدعوم / Unsupported file type', 'danger')
            return redirect(request.url)
    
    return render_template('knowledge/upload.html', lang=lang)

@knowledge_bp.route('/search', methods=['GET', 'POST'])
@login_required
def search():
    user_id = get_jwt_identity()
    user = User.query.get(user_id)
    lang = session.get('language', 'ar')
    
    if request.method == 'POST':
        query = request.form.get('query')
        
        # Get all user documents
        documents = Document.query.filter_by(user_id=user_id).all()
        
        if not documents:
            flash('لا توجد مستندات مرفوعة / No documents uploaded yet', 'warning')
            return redirect(url_for('knowledge.upload'))
        
        # Combine all document texts
        combined_context = "\n\n".join([
            f"مستند: {doc.filename}\n{doc.content_text[:2000]}" 
            for doc in documents if doc.content_text
        ])
        
        if lang == 'ar':
            system_prompt = "أنت مساعد معرفة ذكي. قم بالإجابة على الأسئلة بناءً على المستندات المرفقة فقط."
            user_message = f"""
            السياق من المستندات:
            {combined_context[:4000]}
            
            السؤال: {query}
            
            قم بالإجابة بشكل دقيق مع الإشارة إلى المصدر إن أمكن.
            """
        else:
            system_prompt = "You are an intelligent knowledge assistant. Answer questions based only on the provided documents."
            user_message = f"""
            Context from documents:
            {combined_context[:4000]}
            
            Question: {query}
            
            Provide an accurate answer with source reference when possible.
            """
        
        ai_response = llm_chat(system_prompt, user_message)
        
        # Log AI usage
        ai_log = AILog(
            user_id=user_id,
            module='knowledge',
            prompt=query,
            response=ai_response[:500],
            tokens_used=int(len(ai_response.split()) * 1.3)
        )
        db.session.add(ai_log)
        user.ai_credits_used += int(len(ai_response.split()) * 1.3)
        db.session.commit()
        
        return render_template('knowledge/search_result.html', query=query, answer=ai_response, documents=documents, lang=lang)
    
    return render_template('knowledge/search.html', lang=lang)
